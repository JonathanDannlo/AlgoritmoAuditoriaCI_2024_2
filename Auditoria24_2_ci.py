import os
import math
import pandas as pd
import numpy as np
from docx import Document
from openpyxl import load_workbook
from openpyxl.drawing.image import Image


file_path = '/content/PIAM2024_1.xlsx'
output_pathDoc = '/content/AuditoriaPiam20242Ci.docx'
output_pathXlsx = "/content/AuditoriaPiam20242Ci.xlsx"

columnasValidacionObligatoriedad = [
    'TIPOIDENTIFICACION','IDENTIFICACION','CODIGO','SNIESPROGRAMA','IDMUNICIPIOPROGRAMA','NACIMIENTO','ID_PAIS_NACIMIENTO','IDMUNICIPIONACIMIENTO',
    'ZONARESIDENCIA','ESTRATO','ESTUDIANTEREINGRESO','ANIOINGRESO','PERIODOINGRESO','TELEFONO','CELULAR','EMAILPERSONAL','EMAILINSTITUCIONAL',
    'CREDITOSPENSUM','SEMESTRESPROGRAMA','CREDITOSAPROBADOS','CREDITOSMATRICULADOS','DERECHOS_MATRICULA','SEGURO_ESTUDIANTIL']

valoresValidosTipoId = ['CC', 'DE', 'CE', 'TI', 'PS', 'CA', 'PT']

matriculaBruta = ['DERECHOS_MATRICULA',
                  'BIBLIOTECA_DEPORTES',
                  'LABORATORIOS',
                  'RECURSOS_COMPUTACIONALES',
                  'SEGURO_ESTUDIANTIL',
                  'VRES_COMPLEMENTARIOS',
                  'RESIDENCIAS',
                  'REPETICIONES']
meritoAcademico = ['CONVENIO_DESCENTRALIZACION',
                   'BECA',
                   'MATRICULA_HONOR',
                   'MEDIA_MATRICULA_HONOR',
                   'TRABAJO_GRADO',
                   'DOS_PROGRAMAS',
                   'DESCUENTO_HERMANO',
                   'ESTIMULO_EMP_DTE_PLANTA',
                   'ESTIMULO_CONYUGE',
                   'EXEN_HIJOS_CONYUGE_CATEDRA',
                   'EXEN_HIJOS_CONYUGE_OCASIONAL',
                   'HIJOS_TRABAJADORES_OFICIALES',
                   'ACTIVIDAES_LUDICAS_DEPOR',
                   'DESCUENTOS',
                   'SERVICIOS_RELIQUIDACION',
                   'DESCUENTO_LEY_1171']

def agregar_mensaje(doc, mensaje):
    print(mensaje)
    doc.add_paragraph(mensaje)

def cargar_archivos_y_dataframes(file_path):
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"{file_path} no encontrado.")
    print(f"Archivo {file_path} encontrado.")
    try:
        dic_insumos = pd.read_excel(file_path, sheet_name=['PIAM2024_2_CI', 'ICETEX2024_2', 'CONCI24_1', 'SQ240924'], engine='openpyxl')
        for df in dic_insumos.values():
            df.columns = df.columns.str.strip()
        return dic_insumos['PIAM2024_2_CI'], dic_insumos['ICETEX2024_2'], dic_insumos['SQ240924'], dic_insumos['CONCI24_1']
    except Exception as e:
        raise Exception(f"Error al cargar los DataFrames: {e}")

def obtener_registros_vacios(df, columnas, output_path, doc):
    registros_vacios_total = pd.DataFrame()
    for columna in columnas:
        if columna not in df.columns:
            agregar_mensaje(doc, f"La columna '{columna}' no existe en el DataFrame.")
            continue
        registros_vacios = df[df[columna].isna()].copy()
        num_vacios = len(registros_vacios)
        mensaje = f"Hay {num_vacios} registros vacíos en la columna '{columna}'." if num_vacios > 0 else f"No hay registros vacíos en la columna '{columna}'."
        agregar_mensaje(doc, mensaje)
        if num_vacios > 0:
            registros_vacios['BanderaRegistrosVacios'] = columna
            registros_vacios_total = pd.concat([registros_vacios_total, registros_vacios], ignore_index=True)
    return registros_vacios_total

def validar_tipo_documento(df, doc, valoresValidosTipoId):
    if 'TIPOIDENTIFICACION' in df.columns:
        df_invalidos = df[~df['TIPOIDENTIFICACION'].isin(valoresValidosTipoId)]
        if not df_invalidos.empty:
            mensaje = f"Se encontraron {len(df_invalidos)} registros con tipos de documento no válidos."
            agregar_mensaje(doc, mensaje)
        else:
            agregar_mensaje(doc, "No se encontraron registros con tipos de documento no válidos.")
    else:
        agregar_mensaje(doc, "La columna 'TIPOIDENTIFICACION' no existe en el DataFrame.")
    return df_invalidos if 'TIPOIDENTIFICACION' in df.columns else pd.DataFrame()

def ajustarRegistrosVacios(df, columnas, doc):
    for columna in columnas:
        if columna not in df.columns:
            agregar_mensaje(doc, f"La columna '{columna}' no existe en el DataFrame.")
            continue
        if pd.api.types.is_numeric_dtype(df[columna]):
            df[columna] = df[columna].fillna(0)
            agregar_mensaje(doc, f"Los registros vacíos en la columna '{columna}' han sido rellenados con 0.")
        else:
            df[columna] = df[columna].fillna('NaN')
            agregar_mensaje(doc, f"Los registros vacíos en la columna '{columna}' han sido rellenados con NaN.")
    return df


def verificarInconsistenciasCreditos(df, columna, doc):
    if columna not in df.columns:
        agregar_mensaje(doc, f"La columna '{columna}' no existe en el DataFrame.")
        return pd.DataFrame()
    df['BanderaCreditosRC'] = df[columna] < 15
    df_inconsistenciasRC = df[df['BanderaCreditosRC']]
    agregar_mensaje(doc, f"Se encontraron {len(df_inconsistenciasRC)} programas con inconsistencia en los créditos exigidos por el Registro Calificado.")
    return df_inconsistenciasRC


def verificarInconsistenciasCreditosCantidad(df, creditosRC, creditosAprobados, doc):
    if creditosRC not in df.columns:
        agregar_mensaje(doc, f"La columna '{creditosRC}' no existe en el DataFrame.")
        return pd.DataFrame()
    if creditosAprobados not in df.columns:
        agregar_mensaje(doc, f"La columna '{creditosAprobados}' no existe en el DataFrame.")
        return pd.DataFrame()
    def evaluarInconsistenciaCreditos(row):
        if row[creditosRC] < row[creditosAprobados]:
            return 'Creditos RC menor a los creditos aprobados'
        elif row[creditosRC] == row[creditosAprobados]:
            return 'Creditos RC igual a los creditos aprobados'
        else:
            return None
    df['FlCreditosRCAprobados'] = df.apply(evaluarInconsistenciaCreditos, axis=1)
    df_inconsistenciasRCAprobados = df[df['FlCreditosRCAprobados'].notnull()]
    resumen_inconsistencias = df_inconsistenciasRCAprobados['FlCreditosRCAprobados'].value_counts()
    agregar_mensaje(doc, f"Se encontraron {len(df_inconsistenciasRCAprobados)} inconsistencias en los Créditos del RC y los aprobados:")
    for inconsistencia, cantidad in resumen_inconsistencias.items():
        agregar_mensaje(doc, f"{inconsistencia}: {cantidad} casos")
    return df_inconsistenciasRCAprobados

def ajustarCreditosAprobados(df, creditosRC, creditosAprobados, numSemestres):
    if creditosRC not in df.columns or creditosAprobados not in df.columns or numSemestres not in df.columns:
        print("Una o más columnas especificadas no existen en el DataFrame.")
        return df
    df_ajuste_menor = df[df[creditosRC] < df[creditosAprobados]].copy()
    df_ajuste_igual = df[df[creditosRC] == df[creditosAprobados]].copy()
    df_ajuste_errado = df[df[creditosRC] == 12].copy()
    if not df_ajuste_menor.empty:
        df_ajuste_menor[creditosAprobados] = df_ajuste_menor[creditosRC] - (df_ajuste_menor[creditosRC] / df_ajuste_menor[numSemestres])
    if not df_ajuste_igual.empty:
        df_ajuste_igual[creditosAprobados] = df_ajuste_igual[creditosRC] - (df_ajuste_igual[creditosRC] / df_ajuste_igual[numSemestres])
    if not df_ajuste_errado.empty:
        df_ajuste_errado[creditosRC] = 68
    df.update(df_ajuste_menor)
    df.update(df_ajuste_igual)
    df.update(df_ajuste_errado)
    print(f"Se han ajustado {len(df_ajuste_menor)} registros donde los créditos RC eran menores a los créditos aprobados.")
    print(f"Se han ajustado {len(df_ajuste_igual)} registros donde los créditos RC eran iguales a los créditos aprobados.")
    print(f"Se han ajustado {len(df_ajuste_errado)} registros donde los créditos RC eran iguales a 12.")
    return df

def calcular_matricula(df):
    df['BRUTA'] = df[matriculaBruta].sum(axis=1)
    df['BRUTAORD'] = df['BRUTA'] - df['SEGURO_ESTUDIANTIL']
    df['NETAORD'] = df['BRUTAORD'] - df['VOTO'].abs()
    df['MERITO'] = df[meritoAcademico].sum(axis=1).abs()
    df['MTRNETA'] = df['BRUTA'] - df['VOTO'].abs() - df['MERITO']
    df['NETAAPL'] = df['MTRNETA'] - df['SEGURO_ESTUDIANTIL']
    return df

def validar_matricula_duplicados(df):
    df['FL_NETA'] = df['MTRNETA'] == df['Valor Factura']
    df['DUPLICADO_RECIBO'] = df['RECIBO'].duplicated(keep=False)
    df['DUPLICADO_ID_SNIES'] = df.duplicated(subset=['IDENTIFICACION', 'SNIESPROGRAMA'], keep=False)
    return df

def generarReportePlantillaMatriculados(df):
    columnas_necesarias = ['CREDITOSPENSUM', 'SEMESTRESPROGRAMA', 'VOTO', 'Sublínea Crédito', 'NETAORD', 'MERITO']
    for col in columnas_necesarias:
        if col not in df.columns:
            print(f"Una o más columnas necesarias para el cálculo no existen en el DataFrame: {col}")
            return None, None
    if (df['SEMESTRESPROGRAMA'] == 0).any():
        print("La columna 'SEMESTRES_RC' contiene valores cero, lo que podría causar una división por cero.")
        return None, None
    df['CREDIT_ACAD_A_MATRIC_REGU_SEM'] = ((df['CREDITOSPENSUM'] + 2) / df['SEMESTRESPROGRAMA']).apply(math.ceil)
    df['APOYO_GOB_NAC_DESCUENTO_VOTAC'] = -(df['VOTO']).fillna(0)
    df['APOYO_GOBERNAC_PROGR_PERMANENT'] = 0
    df['APOYO_ALCALDIA_PROGR_PERMANENT'] = 0
    df['DESCUENT_RECURRENTES_DE_LA_IES'] = 0
    df['OTROS_APOYOS_A_LA_MATRICULA'] = 0
    df['OTROS_APOYOS_ADICIONALES'] = 0
    df['DESCUENTOS_ADICIONALES_IES'] = 0
    df.loc[df['Sublínea Crédito'] == '121943 - 121943 SER ESTUDIOSO CUENTA', 'APOYO_ADICIONAL_GOBERNACIONES'] = df['NETAORD']
    df['APOYO_ADICIONAL_ALCALDIAS'] = 0
    df['VAL_NETO_DER_MAT_A_CARGO_EST'] = (
        df['NETAORD'].fillna(0) -
        df['OTROS_APOYOS_ADICIONALES'].fillna(0) -
        df['DESCUENTOS_ADICIONALES_IES'].fillna(0) -
        df['APOYO_ADICIONAL_ALCALDIAS'].fillna(0) -
        df['APOYO_ADICIONAL_GOBERNACIONES'].fillna(0)
    )
    df['VALOR_BRUTO_DERECHOS_COMPLEMEN'] = df['SEGURO_ESTUDIANTIL']
    df['VALOR_NETO_DERECHOS_COMPLEMENT'] = df['SEGURO_ESTUDIANTIL']
    df['CAUSA_NO_ACCESO'] = 0
    columnas_reporte = [
        'TIPOIDENTIFICACION','IDENTIFICACION','CODIGO','SNIESPROGRAMA','IDMUNICIPIOPROGRAMA','NACIMIENTO','ID_PAIS_NACIMIENTO','IDMUNICIPIONACIMIENTO',
        'ZONARESIDENCIA','ESTRATO','ESTUDIANTEREINGRESO','ANIOINGRESO','PERIODOINGRESO', 'NETAORD','TELEFONO','EMAILPERSONAL'
    ]
    columnas_caracterizacion = [
        'TIPOIDENTIFICACION','IDENTIFICACION','SNIESPROGRAMA','IDMUNICIPIOPROGRAMA','CREDITOSPENSUM',
        'CREDITOSAPROBADOS','CREDIT_ACAD_A_MATRIC_REGU_SEM', 'BRUTAORD', 'APOYO_GOB_NAC_DESCUENTO_VOTAC',
        'APOYO_GOBERNAC_PROGR_PERMANENT', 'APOYO_ALCALDIA_PROGR_PERMANENT', 'DESCUENT_RECURRENTES_DE_LA_IES',
        'OTROS_APOYOS_A_LA_MATRICULA', 'NETAORD', 'APOYO_ADICIONAL_GOBERNACIONES', 'APOYO_ADICIONAL_ALCALDIAS',
        'DESCUENTOS_ADICIONALES_IES', 'OTROS_APOYOS_ADICIONALES', 'VAL_NETO_DER_MAT_A_CARGO_EST',
        'VALOR_BRUTO_DERECHOS_COMPLEMEN', 'VALOR_NETO_DERECHOS_COMPLEMENT', 'CAUSA_NO_ACCESO'
    ]
    for columna in columnas_reporte:
        if columna not in df.columns:
            print(f"La columna '{columna}' no existe en el DataFrame.")
            return None, None
    df_reporte = df[columnas_reporte].copy()
    for columna in columnas_caracterizacion:
        if columna not in df.columns:
            print(f"La columna '{columna}' no existe en el DataFrame.")
            return None, None
    df_reporte1 = df[columnas_caracterizacion].copy()
    if 'NACIMIENTO' in df_reporte1.columns:
        df_reporte1['NACIMIENTO'] = pd.to_datetime(df_reporte1['FECHA_NACIMIENTO']).dt.strftime('%Y/%m/%d')
    print("Reporte 'PlantillaMatriculados' generado exitosamente.")
    print("Reporte 'PlantillaCaracterizacion' generado exitosamente.")
    return df_reporte, df_reporte1

file_path = '/content/PIAM2024_1.xlsx'
output_pathDoc = '/content/AuditoriaPiam20242Ci.docx'
output_pathXlsx = '/content/AuditoriaPiam20242Ci.xlsx'

doc = Document()
piam20242, icetex, facturacion20242, conci241 = cargar_archivos_y_dataframes(file_path)
registros_vacios_total = obtener_registros_vacios(piam20242, columnasValidacionObligatoriedad, output_pathDoc, doc)
df_resultado0 = ajustarRegistrosVacios(piam20242, columnasValidacionObligatoriedad, doc)
df_resultado1 = validar_tipo_documento(df_resultado0, doc, valoresValidosTipoId)
df_inconsistencias = verificarInconsistenciasCreditos(df_resultado0, 'CREDITOSPENSUM', doc)
df_inconsistenciasCantidad = verificarInconsistenciasCreditosCantidad(df_resultado0, 'CREDITOSPENSUM', 'CREDITOSAPROBADOS', doc)
df_ajustado = ajustarCreditosAprobados(df_resultado0,'CREDITOSPENSUM','CREDITOSAPROBADOS','SEMESTRESPROGRAMA')
df_inconsistenciasCantidad1 = verificarInconsistenciasCreditosCantidad(df_ajustado, 'CREDITOSPENSUM', 'CREDITOSAPROBADOS', doc)
df_ajustado1 = calcular_matricula(df_ajustado)
doc.save(output_pathDoc)


df_ajustado1['CODIGO'] = df_ajustado1['CODIGO'].astype(str)
df_ajustado1['RECIBO'] = df_ajustado1['RECIBO'].astype(str)
icetex['Código'] = icetex['Código'].astype(str)
conci241['codigo'] = conci241['codigo'].astype(str)
facturacion20242['Documento'] = facturacion20242['Documento'].astype(str)

df_piam20242_ajustado = pd.merge(
    df_ajustado1, icetex[['Código', 'Sublínea Crédito','Relación de Giro','Total a Girar']],
    left_on='CODIGO',
    right_on='Código',
    how='left')
df_piam20242_ajustado['Coincidencia_ICETEX'] = df_piam20242_ajustado['Código'].notna()

df_piam20242_ajustado1 = pd.merge(
    df_piam20242_ajustado,
    conci241[['codigo', 'ESTADO_GIRO','PERIODOS_A_FINANCIAR']],
    left_on='CODIGO',
    right_on='codigo',
    how='left')
df_piam20242_ajustado1['Coincidencia_CONCI'] = df_piam20242_ajustado1['codigo'].notna()

df_piam20242_ajustado2 = pd.merge(
    df_piam20242_ajustado1,
    facturacion20242,
    left_on='RECIBO',
    right_on='Documento',
    how='left')
df_piam20242_ajustado2['Coincidencia_FACTURACION'] = df_piam20242_ajustado2['Documento'].notna()


df_facturacion_sin_coincidencia = pd.merge(
    df_piam20242_ajustado1,
    facturacion20242,
    left_on='RECIBO',
    right_on='Documento',
    how='right',
    indicator=True
)

total_registros = len(df_piam20242_ajustado2)
coincidencias_icetex = df_piam20242_ajustado2['Coincidencia_ICETEX'].sum()
coincidencias_conci = df_piam20242_ajustado2['Coincidencia_CONCI'].sum()
coincidencias_facturacion = df_piam20242_ajustado2['Coincidencia_FACTURACION'].sum()
sin_coincidencia_icetex = total_registros - coincidencias_icetex
sin_coincidencia_conci = total_registros - coincidencias_conci
sin_coincidencia_facturacion = total_registros - coincidencias_facturacion

resumen = f"""
Resumen de Coincidencias:
Total de registros: {total_registros}
1. Coincidencias con ICETEX:
   - Registros con coincidencia: {coincidencias_icetex}
   - Registros sin coincidencia: {sin_coincidencia_icetex}
2. Coincidencias con CONCI:
   - Registros con coincidencia: {coincidencias_conci}
   - Registros sin coincidencia: {sin_coincidencia_conci}
3. Coincidencias con FACTURACIÓN:
   - Registros con coincidencia: {coincidencias_facturacion}
   - Registros sin coincidencia: {sin_coincidencia_facturacion}
"""
doc.add_heading('Resumen de Coincidencias', level=1)
doc.add_paragraph(resumen)
doc.save(output_pathDoc)


registros_no_en_ajustado1 = df_facturacion_sin_coincidencia[df_facturacion_sin_coincidencia['_merge'] == 'right_only']
total_no_en_ajustado1 = len(registros_no_en_ajustado1)
resumen_no_coincidencia = f"""
Registros en 'facturacion20242' pero no en 'ajustado1':
Total de registros sin coincidencia: {total_no_en_ajustado1}
"""
doc.add_heading('Registros en Facturación no coincidentes con Ajustado1', level=2)
doc.add_paragraph(resumen_no_coincidencia)
doc.save(output_pathDoc)

df_piam20242_ajustado2_validado = validar_matricula_duplicados(df_piam20242_ajustado2)

plantilla_matriculados, plantilla_caracterizacion = generarReportePlantillaMatriculados(df_piam20242_ajustado2_validado)

with pd.ExcelWriter(output_pathXlsx, engine='xlsxwriter') as writer:
    """df_resultado0.to_excel(writer, sheet_name='PIAM2024_2_CI', index=False)
    registros_vacios_total.to_excel(writer, sheet_name='RegistrosVacios', index=False)
    if not df_resultado1.empty:
        df_resultado1.to_excel(writer, sheet_name='TipoDocumentoInvalido', index=False)
    if not df_inconsistencias.empty:
        df_inconsistencias.to_excel(writer, sheet_name='InconsistenciasCreditosRC', index=False)
    if not df_inconsistenciasCantidad.empty:
        df_inconsistenciasCantidad.to_excel(writer, sheet_name='InconsistenciasCreditosCantidad', index=False)
    if not df_ajustado.empty:
        df_ajustado.to_excel(writer, sheet_name='PIAM20242CIA', index=False)
    if not df_ajustado1.empty:
        df_ajustado1.to_excel(writer, sheet_name='PIAM20242CI_AM', index=False)
    if not df_piam20242_ajustado.empty:
        df_piam20242_ajustado.to_excel(writer, sheet_name='PIAM20242_AJ', index=False)
    if not df_piam20242_ajustado1.empty:
        df_piam20242_ajustado1.to_excel(writer, sheet_name='PIAM20242_AJ1', index=False)
    if not df_piam20242_ajustado2.empty:
        df_piam20242_ajustado2.to_excel(writer, sheet_name='PIAM20242_AJ2', index=False)"""
    if not registros_no_en_ajustado1.empty:
        registros_no_en_ajustado1.to_excel(writer, sheet_name='NoCoincidentes_Facturacion', index=False)
    if not df_piam20242_ajustado2_validado.empty:
        df_piam20242_ajustado2_validado.to_excel(writer, sheet_name='PIAM20242_AJ2V', index=False)
    if plantilla_matriculados is not None:
      plantilla_matriculados.to_excel(writer, sheet_name='PlantillaMatriculados', index=False)
      print('Se ha generado la plantilla de matriculados')
    if plantilla_caracterizacion is not None:
      plantilla_caracterizacion.to_excel(writer, sheet_name='PlantillaCaracterizacion', index=False)
      print('Se ha generado la plantilla de caracterizacion')

print("Los resultados han sido guardados en el documento y archivo Excel.")
